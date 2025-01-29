import os
import tempfile
import pythoncom  # Importujemy pythoncom, aby użyć CoInitialize
import pandas as pd
from flask import Flask, request, send_file, render_template_string
from openpyxl import load_workbook
import re
import docx
import win32com.client as win32

app = Flask(__name__)

TEMPLATE_PATH = 'T_Fead-110414.xlsx'
LOGO_PATH = '/static/logo.png'
LIST_MATERIAL_PATH = 'list material.xls'

HTML_FORM = """
<!DOCTYPE html>
<html lang="nl">
<head>
    <meta charset="UTF-8">
    <title>Excel Rapport Generator</title>
    <style>
        body {
            display: flex;
            justify-content: center;
            align-items: center;
            min-height: 100vh;
            margin: 0;
            font-family: Arial, sans-serif;
            flex-direction: column;
            background-color: #ffffff;
        }
        .header {
            text-align: center;
            margin-bottom: 30px;
        }
        .logo {
            width: 800px;
            height: auto;
            margin-bottom: 20px;
        }
        .tool-report {
            font-size: 48px;
            font-weight: bold;
            color: #333;
            margin-bottom: 40px;
        }
        .form-container {
            width: 90%;
            max-width: 600px;
            border: 2px solid #808080;
            padding: 20px;
            border-radius: 10px;
            background-color: #e0e0e0;
        }
        .form-container label,
        .form-container input {
            display: block;
            margin: 10px 0;
            width: 100%;
        }
        button {
            background-color: #808080;
            color: white;
            border: none;
            padding: 10px 20px;
            font-size: 16px;
            cursor: pointer;
            border-radius: 5px;
            margin-top: 20px;
            width: 100%;
        }
        button:hover {
            background-color: #A9A9A9;
        }
    </style>
</head>
<body>
    <div class="header">
        <img src="{{ logo_path }}" alt="Logo" class="logo">
        <div class="tool-report">TOOL REPORT</div>
    </div>
    <div class="form-container">
        <form action="/generate" method="post">
            <label for="tool_number">Gereedschapsnummer (8 cijfers):</label>
            <input type="text" id="tool_number" name="tool_number" required pattern="\\d{8}"><br>

            <label for="directory_path">Waar is het pad naar de map:</label>
            <input type="text" id="directory_path" name="directory_path" required><br>

            <button type="submit">Genereer TOOL REPORT</button>
        </form>
    </div>
</body>
</html>
"""

# Funkcja do konwersji pliku .doc na .docx w katalogu tymczasowym
def convert_doc_to_docx_temp(file_path):
    try:
        pythoncom.CoInitialize()  # Inicjalizacja COM
        word = win32.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(file_path)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx:
            temp_docx_path = temp_docx.name
        doc.SaveAs(temp_docx_path, FileFormat=16)  # 16 to format .docx
        doc.Close()
        word.Quit()
        print(f"Przekonwertowano plik .doc na .docx: {temp_docx_path}")
        return temp_docx_path
    except Exception as e:
        print(f"Błąd podczas konwersji pliku .doc na .docx: {e}")
        return None
    finally:
        pythoncom.CoUninitialize()  # Dekonstrukcja COM

# Funkcja do przeszukiwania plików GB w określonym katalogu
def extract_bestelnummer_from_gb_files(directory_path, tool_number):
    try:
        # Przeszukiwanie określonego katalogu
        if os.path.exists(directory_path):
            for filename in os.listdir(directory_path):
                file_path = os.path.join(directory_path, filename)
                if filename.startswith("GB-") and filename.endswith(".docx"):
                    print(f"Przetwarzanie pliku .docx: {file_path}")
                    doc = docx.Document(file_path)
                    tables = doc.tables
                    if len(tables) < 1:
                        print(f"Brak tabel w pliku .docx: {file_path}")
                        continue
                    for table in tables:
                        for row in table.rows:
                            tool_cell = row.cells[0].text.strip()
                            description_cell = row.cells[2].text.strip() if len(row.cells) > 2 else ""
                            print(f"Sprawdzanie wiersza: tool_cell='{tool_cell}', description_cell='{description_cell}'")
                            if tool_cell == tool_number:
                                barcode_match = re.findall(r'\*(\d+)\*', description_cell)
                                if barcode_match:
                                    return f"*{barcode_match[-1]}*"
                elif filename.startswith("GB-") and filename.endswith(".doc"):
                    print(f"Przetwarzanie pliku .doc: {file_path}")
                    docx_path = convert_doc_to_docx_temp(file_path)
                    if docx_path:
                        doc = docx.Document(docx_path)
                        tables = doc.tables
                        if len(tables) < 1:
                            print(f"Brak tabel w przekonwertowanym pliku .docx: {docx_path}")
                            continue
                        for table in tables:
                            for row in table.rows:
                                tool_cell = row.cells[0].text.strip()
                                description_cell = row.cells[2].text.strip() if len(row.cells) > 2 else ""
                                print(f"Sprawdzanie wiersza w przekonwertowanym pliku: tool_cell='{tool_cell}', description_cell='{description_cell}'")
                                if tool_cell == tool_number:
                                    barcode_match = re.findall(r'\*(\d+)\*', description_cell)
                                    if barcode_match:
                                        return f"*{barcode_match[-1]}*"
                        # Usunięcie pliku .docx po przetworzeniu
                        try:
                            os.remove(docx_path)
                            print(f"Usunięto plik .docx: {docx_path}")
                        except Exception as e:
                            print(f"Błąd podczas usuwania pliku .docx: {e}")
        else:
            print(f"Katalog nie istnieje: {directory_path}")
        return "Nie znaleziono"
    except Exception as e:
        print(f"Błąd podczas odczytu plików GB: {e}")
        return "Nie znaleziono"

@app.route('/')
def home():
    return render_template_string(HTML_FORM, logo_path=LOGO_PATH)

@app.route('/generate', methods=['POST'])
def generate_excel():
    tool_number = request.form.get('tool_number', '00000000')
    directory_path = request.form.get('directory_path', 'N:\\Werkvoorbereiding\\Klantenbestand\\Haton\\HAB050003328\\Documenten\\Variaxis 630')
    tool_bestelnummer = extract_bestelnummer_from_gb_files(directory_path, tool_number)

    # Odczyt pliku z listą materiałów, aby uzyskać nazwę narzędzia na podstawie numeru narzędzia
    try:
        df = pd.read_excel(LIST_MATERIAL_PATH, header=None)

        # Upewnij się, że porównywane są stringi
        df[0] = df[0].astype(str)  # Konwersja pierwszej kolumny (Tool Number) na string

        # Szukaj numeru narzędzia w kolumnie A
        tool_data = df[df[0] == tool_number]

        if not tool_data.empty:
            tool_name = tool_data.iloc[0, 2]  # Pobierz nazwę narzędzia z kolumny C
        else:
            tool_name = f"Gereedschapsnummer {tool_number} niet gevonden in materiaallijst."
            print(tool_name)
            return render_template_string(HTML_FORM, logo_path=LOGO_PATH)
    except Exception as e:
        print(f"Fout bij lezen van materiaallijst: {e}")
        return f"Fout bij lezen van materiaallijst: {e}", 500

    # Upewnij się, że mamy poprawną nazwę narzędzia i wyciągnij wszystko po "OH"
    holder_name = ""
    if tool_name and "Gereedschapsnummer" not in tool_name:
        oh_match = re.search(r'OH\d+-(.*)', tool_name)
        if oh_match:
            extracted_name = oh_match.group(1)  # Wyciągnij tekst po "OHXX-"
            holder_name = f"Houder: BT40-{extracted_name}"  # Dodaj "Houder: BT40-" na początku
            print(f"Houder Naam: {holder_name}")
        else:
            print("Patroon niet gevonden in Tool Name.")
    else:
        print("Tool Name is ongeldig of niet gevonden.")

    output_path = f"{tool_number}.xlsx"

    # Wyciągnij wartości OH, LPR, Z i DC z nazwy narzędzia za pomocą regex
    oh_match = re.search(r'OH(\d+)', tool_name)
    oh_value = int(oh_match.group(1)) if oh_match else 0
    lpr_match = re.search(r'LPR(\d+)', tool_name)
    lpr_value = int(lpr_match.group(1)) if lpr_match else 0
    z_match = re.search(r'Z(\d+)', tool_name)
    z_value = int(z_match.group(1)) if z_match else 0
    dc_match = re.search(r'DC(\d+)', tool_name)
    dc_value = int(dc_match.group(1)) if dc_match else 0
    gl_value = oh_value + lpr_value

    print(f"OH Waarde: {oh_value}, LPR Waarde: {lpr_value}, Z Waarde: {z_value}, DC Waarde: {dc_value}, GL Waarde: {gl_value}")

    # Załaduj szablon Excel
    try:
        workbook = load_workbook(TEMPLATE_PATH)
        sheet = workbook.active  # Zakładamy, że dane trafiają do pierwszego arkusza

        # Zapisz dynamiczne wartości w pliku Excel
        sheet.cell(row=6, column=25, value=f"OH={oh_value}")  # Zapisz OH w Y6
        sheet.cell(row=5, column=25, value=f"GL={gl_value}")  # Zapisz GL=OH+LPR w Y5
        sheet.cell(row=6, column=40, value=f"D={dc_value}")  # Zapisz DC w AN6
        sheet.cell(row=7, column=40, value=f"Z={z_value}")  # Zapisz Z w AN7

        # Zapisz nazwę Holder w P16
        print(f"Zapisz nazwę Holder w P16: {holder_name}")
        sheet.cell(row=16, column=16, value=holder_name)  # Zapisz holder_name w P16

        # Zapisz pozostałe wartości w arkuszu Excel
        sheet.cell(row=2, column=2, value=tool_number)  # Zapisz numer narzędzia w B2
        sheet.cell(row=14, column=16, value=f"Tool: {tool_name}")  # Zapisz nazwę narzędzia w P14
        sheet.cell(row=14, column=2, value=tool_bestelnummer)  # Zapisz zamówienie narzędzia w B14

        # Zapisz zaktualizowany plik
        workbook.save(output_path)
        workbook.close()

    except Exception as e:
        print(f"Fout bij bijwerken van Excel: {e}")
        return f"Fout bij bijwerken van Excel: {e}", 500

    # Wyślij plik do pobrania
    return send_file(output_path, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True, port=5000)